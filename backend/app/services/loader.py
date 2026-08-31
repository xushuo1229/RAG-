"""文档解析与切分：支持 txt / md / pdf / docx。"""
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter

SUPPORTED = {"txt", "md", "pdf", "docx"}


def read_text(path: Path) -> str:
    """读取纯文本，优先 UTF-8，失败回退 GBK。"""
    raw = path.read_bytes()
    for enc in ("utf-8", "gbk"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def parse_content(path: Path, filename: str) -> str:
    """根据扩展名解析文档为纯文本。"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("txt", "md"):
        return read_text(path)
    if ext == "pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if ext == "docx":
        from docx import Document

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    raise ValueError(f"不支持的文件类型：{ext or '未知'}")


def split_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """按中文友好分隔符递归切分，返回清洗后的非空文本块。"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
    )
    chunks = splitter.split_text(text)
    return [c.strip() for c in chunks if c.strip()]